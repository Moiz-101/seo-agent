import os
import re

from docx import Document

from config import DATA_DIR

DOCS_DIR = os.path.join(DATA_DIR, "generated_docs")
os.makedirs(DOCS_DIR, exist_ok=True)


def markdown_to_docx(markdown_text: str, filename: str) -> str:
    """Very small markdown->docx converter: handles #/##/### headings,
    - bullet lists, and plain paragraphs. Enough for LLM-generated articles.
    """
    doc = Document()

    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            doc.add_heading(heading_match.group(2), level=level)
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue

        # strip stray markdown bold/italic markers for cleanliness
        clean = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
        clean = re.sub(r"\*(.*?)\*", r"\1", clean)
        doc.add_paragraph(clean)

    path = os.path.join(DOCS_DIR, filename)
    doc.save(path)
    return path


def save_article_as_docx(article_markdown: str, primary_keyword: str) -> str:
    safe_name = re.sub(r"[^a-z0-9]+", "-", primary_keyword.lower()).strip("-")
    filename = f"{safe_name}.docx"
    return markdown_to_docx(article_markdown, filename)
